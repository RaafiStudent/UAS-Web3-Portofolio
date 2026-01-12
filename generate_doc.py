
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =================KONFIGURASI=================
OUTPUT_FILENAME = "DOKUMENTASI_SOURCE_CODE_PROJECT.docx"

# Folder yang AKAN DILEWATI (Filter Sampah)
EXCLUDE_DIRS = {
    'node_modules', 'vendor', '.git', '.idea', '.vscode', 
    'build', 'dist', 'storage', 'tmp', '__pycache__', 
    'bin', 'obj', 'debug', 'release', 'gradle', '.gradle',
    'ios/Flutter', 'android/app/build'
}

# Ekstensi file yang AKAN DIAMBIL (Multi-Language)
INCLUDE_EXTENSIONS = {
    # Web / Frontend
    '.html', '.css', '.scss', '.js', '.jsx', '.ts', '.tsx', '.vue', '.svelte', '.json',
    # Backend
    '.php', '.py', '.go', '.java', '.rb', '.sql', '.cs',
    # Mobile
    '.dart', '.kt', '.swift', '.xml',
    # Config / Env / Others
    '.yaml', '.yml', '.gradle', '.properties', '.env.example', 'podfile', 'dockerfile'
}

# File spesifik yang WAJIB diambil (meskipun ekstensinya mungkin tidak umum)
IMPORTANT_FILES = {
    'package.json', 'composer.json', 'pubspec.yaml', 'requirements.txt',
    'build.gradle', 'Podfile', 'Dockerfile', 'docker-compose.yml', '.env.example'
}
# =============================================

def create_document():
    doc = Document()
    
    # --- Halaman Judul ---
    title = doc.add_heading('DOKUMENTASI SOURCE CODE\n& DEPENDENCIES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Mendapatkan path root tempat script dijalankan
    root_path = os.getcwd()
    print(f"[*] Memulai scanning di: {root_path}")

    file_count = 0

    # Mulai menelusuri folder (Walk)
    for dirpath, dirnames, filenames in os.walk(root_path):
        # Filter Folder Sampah secara in-place agar os.walk tidak masuk ke dalamnya
        dirnames[:] = [d for d in dirnames if d not in EXCLUDE_DIRS]

        for filename in filenames:
            file_path = os.path.join(dirpath, filename)
            _, ext = os.path.splitext(filename)
            
            # Cek apakah file harus diambil
            is_extension_valid = ext.lower() in INCLUDE_EXTENSIONS
            is_important_file = filename in IMPORTANT_FILES
            
            if is_extension_valid or is_important_file:
                # Skip file script ini sendiri & file output docx
                if filename == os.path.basename(__file__) or filename == OUTPUT_FILENAME:
                    continue

                try:
                    # Buat Path Relatif untuk Header (contoh: src/Controllers/UserController.php)
                    relative_path = os.path.relpath(file_path, root_path)
                    print(f"    -> Memproses: {relative_path}")

                    # Tambahkan Header Nama File
                    heading = doc.add_heading(relative_path, level=2)
                    heading.style.font.color.rgb = RGBColor(0, 51, 102) # Biru Gelap

                    # Baca isi file
                    content = ""
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    except UnicodeDecodeError:
                        # Fallback jika bukan UTF-8 (misal file windows lama)
                        try:
                            with open(file_path, 'r', encoding='latin-1') as f:
                                content = f.read()
                        except:
                            content = "[ERROR: File encoding tidak didukung atau file binary]"

                    # Tulis isi kodingan
                    paragraph = doc.add_paragraph(content)
                    run = paragraph.runs[0]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    
                    # Tambahkan spasi antar file
                    doc.add_paragraph("-" * 50) 
                    file_count += 1

                except Exception as e:
                    print(f"[!] Gagal memproses {filename}: {e}")

    # Simpan hasil
    try:
        doc.save(OUTPUT_FILENAME)
        print(f"\n[SUCCESS] Dokumentasi selesai! Total file: {file_count}")
        print(f"File tersimpan sebagai: {OUTPUT_FILENAME}")
    except PermissionError:
        print(f"\n[ERROR] Gagal menyimpan. Tutup file {OUTPUT_FILENAME} jika sedang dibuka!")

if __name__ == "__main__":
    create_document()

